"""
Build PROJECT_REPORT.docx from README.md.

Renders the Markdown source as a formatted Word document with real Word tables,
the curated figures from outputs/report_figures/ embedded with captions, and an
auto-updating table of contents.

Run tools/prepare_report_figures.py first: it selects the figures the report uses,
strips the pipeline's baked-in "Figure N" titles, and renames them into report
order, so the caption is the only place a figure number appears.

    python tools/prepare_report_figures.py
    python tools/build_report_docx.py

Requires python-docx and Pillow.
"""

from __future__ import annotations

import re
import struct
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "README.md"
# Curated, title-stripped, report-ordered figures — see tools/prepare_report_figures.py.
FIGDIR = ROOT / "outputs" / "report_figures"
OUT = ROOT / "PROJECT_REPORT.docx"

BODY_FONT = "Cambria"
BODY_SIZE = 10.5   # every run of regular prose uses this: paragraphs, lists,
                   # block quotations and the abstract
HEAD_FONT = "Calibri Light"
MONO_FONT = "Consolas"
TABLE_FONT = "Calibri"

INK = RGBColor(0x1A, 0x1A, 0x1A)
ACCENT = RGBColor(0x1F, 0x3B, 0x57)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)

HEADER_FILL = "1F3B57"
ZEBRA_FILL = "F2F5F8"
QUOTE_FILL = "F4F1EA"
CODE_FILL = "F3F3F3"

CONTENT_WIDTH_IN = 6.27  # A4 (8.27in) minus 1in margins each side

# Figure captions live on their own paragraph style so that the list of figures
# and the list of tables can be built from separate TOC fields.
FIGURE_CAPTION_STYLE = "Figure Caption"
TOC_FIGURES = r'TOC \h \z \t "' + FIGURE_CAPTION_STYLE + r',1"'
TOC_TABLES = r'TOC \h \z \t "Caption,1"'

# ---------------------------------------------------------------------------
# Curated captions, in document order. A mismatch between this list and the number
# of tables found in README.md aborts the build, rather than silently
# shifting every caption by one.
# ---------------------------------------------------------------------------
TABLE_CAPTIONS = [
    "How this project differs from prior work",
    "Survey collection summary",
    "The twelve Likert items",
    "Sample composition",
    "Reliability and factorability of the twelve-item pool",
    "Factor structure and the resulting subscales",
    "Clustering models and their settings",
    "The eight cluster-count criteria and the k each preferred",
    "Clustering algorithms compared at k = 3",
    "Feature sets compared (k-means, k = 3)",
    "Agreement between the WEKA runs and the notebook",
    "Decision tree comparison — CART against J48",
    "Persona profiles, as z-scores",
    "Share of variance explained by each dimension",
    "Cluster validity and stability checks",
    "Association between the personas and the held-out variables",
    "Most common free-text themes",
    "Work distribution among teammates",
]

# Figures are emitted at the end of the section whose heading is keyed here, and
# must be listed in the same order as prepare_report_figures.py numbers them: the
# caption number is read from the filename, so the two have to agree.
FIGURES = {
    "1.2 Structure": [
        ("figure_01_likert_distribution.png",
         "Response distribution per item on the raw scale. Items marked (R) are positively worded "
         "and are reverse-coded downstream (§2.2)."),
    ],
    "2.4 Reliability check": [
        ("figure_02_interitem_correlations.png",
         "Inter-item correlations, with related items aligned. Blocks of related items are "
         "visible, but no single dominant factor."),
    ],
    "2.5 Factor analysis and subscale construction": [
        ("figure_03_efa_loadings.png",
         "The four-factor varimax solution. Bold marks a loading of 0.30 or above; each item marks "
         "exactly one facet, with no cross-loading."),
    ],
    "3.3 Choosing the number of clusters": [
        ("figure_04_validity_indices.png",
         "Validity indices by k for k-means, across every candidate feature space."),
        ("figure_05_gap_statistic.png",
         "Gap statistic on the chosen feature set over 50 uniform reference draws — the one "
         "criterion that can vote for k = 1, and it does."),
        ("figure_06_bootstrap_stability.png",
         "Reproducibility of the partition under resampling: 200 subsamples at 80% of n per k. The "
         "0.75 line is Hennig's threshold for a valid, stable cluster (§0.5)."),
    ],
    "4.1 Clustering algorithms compared (k = 3, final feature set)": [
        ("figure_07_consensus_matrix.png",
         "Consensus matrix at k = 3, students ordered by cluster."),
    ],
    "4.4 Decision tree — describing the personas as rules": [
        ("figure_08_decision_tree.png",
         "Persona membership as a rule set, to depth 4. Interpretability, not prediction: the "
         "labels it is trained on came from the clusterer."),
    ],
    "5.1 Profiles": [
        ("figure_09_persona_profiles.png",
         "Persona profiles as z-scores. Read across a row for a shape, down a column for who is "
         "highest."),
        ("figure_10_persona_cards.png",
         "Persona cards, each with its suggested intervention."),
    ],
    "5.3 Are the groups real?": [
        ("figure_11_holdout_centroids.png",
         "Centroid reproduction on the 30% holdout."),
    ],
    "6.2 Stressors the survey never asked about": [
        ("figure_12_unasked_stressors.png",
         "What only the free text can deliver: stressors students volunteered that no Likert item "
         "covers."),
    ],
}


# ---------------------------------------------------------------------------
# low-level docx helpers
# ---------------------------------------------------------------------------
def shade(cell_or_par, fill: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), fill)
    target = cell_or_par._tc.get_or_add_tcPr() if hasattr(cell_or_par, "_tc") else cell_or_par._p.get_or_add_pPr()
    target.append(el)


def set_repeat_header(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def set_autofit(table) -> None:
    """Fill the text column, but let Word distribute width by cell content.

    Without this, Word gives every column an equal share, which badly over-wraps
    the label column of wide tables whose remaining columns hold short numbers.
    """
    tblPr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblLayout"):
        for old in tblPr.findall(qn(tag)):
            tblPr.remove(old)
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), "5000")
    w.set(qn("w:type"), "pct")
    tblPr.append(w)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "autofit")
    tblPr.append(layout)


def set_cell_margins(table, top=60, bottom=60, left=100, right=100) -> None:
    tblPr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for tag, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tblPr.append(mar)


def add_hyperlink(paragraph, url: str, text: str, bold=False, italic=False):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    for tag in ("w:u",):
        node = OxmlElement(tag)
        node.set(qn("w:val"), "single")
        rPr.append(node)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F3B57")
    rPr.append(color)
    if bold:
        rPr.append(OxmlElement("w:b"))
    if italic:
        rPr.append(OxmlElement("w:i"))
    run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def add_toc_field(paragraph, instruction: str = r'TOC \o "1-3" \h \z \u') -> None:
    r = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose Update Field to build the table of contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for node in (fld_begin, instr, fld_sep, placeholder, fld_end):
        r._r.append(node)


def force_field_update_on_open(document) -> None:
    settings = document.settings.element
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    settings.append(el)


def png_size(path: Path) -> tuple[int, int]:
    data = path.read_bytes()[:33]
    return struct.unpack(">II", data[16:24])


# ---------------------------------------------------------------------------
# inline markdown
# ---------------------------------------------------------------------------
INLINE = re.compile(
    r"(\*\*.+?\*\*"           # bold
    r"|\*[^*\n]+?\*"          # italic
    r"|`[^`]+?`"              # code
    r"|\[[^\]]+?\]\([^)]+?\)" # link
    r"|https?://[^\s)]+)"     # bare url
)


def emit_inline(paragraph, text: str, base_bold=False, base_italic=False, size=None) -> None:
    """Render a Markdown inline string into runs on `paragraph`."""
    text = text.replace(" ", " ")
    for piece in INLINE.split(text):
        if not piece:
            continue
        bold, italic, code, url, shown = base_bold, base_italic, False, None, piece

        if piece.startswith("**") and piece.endswith("**") and len(piece) > 4:
            bold, shown = True, piece[2:-2]
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            italic, shown = True, piece[1:-1]
        elif piece.startswith("`") and piece.endswith("`") and len(piece) > 2:
            code, shown = True, piece[1:-1]
        elif piece.startswith("[") :
            m = re.match(r"\[([^\]]+?)\]\(([^)]+?)\)", piece)
            if m:
                shown, target = m.group(1), m.group(2)
                # Internal anchors carry no meaning in Word; render as plain text.
                if target.startswith(("http://", "https://")):
                    url = target
                shown = re.sub(r"^`|`$", "", shown)
        elif piece.startswith("http"):
            url, shown = piece, piece

        if url:
            add_hyperlink(paragraph, url, shown, bold=bold, italic=italic)
            continue

        run = paragraph.add_run(shown)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = INK
        if size:
            run.font.size = size
        if code:
            run.font.name = MONO_FONT
            run.font.size = Pt((size or Pt(BODY_SIZE)).pt - 1)


def strip_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+?)\]\([^)]+?\)", r"\1", text)
    return text.replace("**", "").replace("`", "").replace("*", "")


# ---------------------------------------------------------------------------
# document scaffolding
# ---------------------------------------------------------------------------
def build_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)
    normal.font.color.rgb = INK
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = normal.paragraph_format
    pf.space_after = Pt(7)
    pf.line_spacing = 1.15
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    sizes = {"Heading 1": 19, "Heading 2": 14, "Heading 3": 11.5}
    for name, size in sizes.items():
        st = doc.styles[name]
        st.font.name = HEAD_FONT
        st.font.size = Pt(size)
        st.font.bold = name != "Heading 1"
        st.font.color.rgb = ACCENT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), HEAD_FONT)
        st.paragraph_format.space_before = Pt(16 if name == "Heading 1" else 12)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Word invents TOC 1/2/3 when it populates a field; predefining them keeps the
    # contents and the figure/table list compact instead of running onto a stray page.
    for name, size, indent in (("TOC 1", 9.5, 0.0), ("TOC 2", 9.5, 0.22), ("TOC 3", 9.5, 0.44)):
        try:
            st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            st = doc.styles[name]
        st.font.name = TABLE_FONT
        st.font.size = Pt(size)
        st.font.color.rgb = INK
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after = Pt(0)
        st.paragraph_format.line_spacing = 1.0
        st.paragraph_format.left_indent = Inches(indent)

    # Table captions keep the built-in "Caption" style; figure captions get their
    # own, so the list of figures and the list of tables can each collect one of
    # them with a TOC field.
    for name, align in (("Caption", WD_ALIGN_PARAGRAPH.LEFT),
                        (FIGURE_CAPTION_STYLE, WD_ALIGN_PARAGRAPH.CENTER)):
        try:
            cap = doc.styles[name]
        except KeyError:
            cap = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        cap.font.name = TABLE_FONT
        cap.font.size = Pt(8.5)
        cap.font.bold = False
        cap.font.italic = False
        cap.font.color.rgb = MUTED
        cap.paragraph_format.space_before = Pt(3)
        cap.paragraph_format.space_after = Pt(12)
        cap.paragraph_format.alignment = align


def page_setup(doc: Document) -> None:
    s = doc.sections[0]
    s.page_width = Inches(8.27)
    s.page_height = Inches(11.69)
    for attr in ("left_margin", "right_margin"):
        setattr(s, attr, Inches(1.0))
    s.top_margin = Inches(0.9)
    s.bottom_margin = Inches(0.9)

    footer = s.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    for tag, txt in (("begin", None), (None, "PAGE"), ("separate", None), (None, "1"), ("end", None)):
        if tag:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), tag)
        elif txt == "PAGE":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = " PAGE "
        else:
            el = OxmlElement("w:t")
            el.text = txt
        run._r.append(el)
    run.font.name = TABLE_FONT
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED


def title_page(doc: Document, title: str, meta: list[str], abstract: list[str]) -> None:
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run(title)
    run.font.name = HEAD_FONT
    run.font.size = Pt(25)
    run.font.color.rgb = ACCENT
    run.bold = False

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = rule.add_run("—" * 18)
    rr.font.color.rgb = RGBColor(0xB0, 0xB8, 0xC0)
    rule.paragraph_format.space_after = Pt(22)

    for line in meta:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mp.paragraph_format.space_after = Pt(4)
        emit_inline(mp, line, size=Pt(11.5))

    for _ in range(2):
        doc.add_paragraph()

    if abstract:
        hp = doc.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run("ABSTRACT")
        hr.font.name = HEAD_FONT
        hr.font.size = Pt(10)
        hr.font.color.rgb = ACCENT
        hr.bold = True
        hp.paragraph_format.space_after = Pt(8)

        ap = doc.add_paragraph()
        ap.paragraph_format.left_indent = Inches(0.55)
        ap.paragraph_format.right_indent = Inches(0.55)
        ap.paragraph_format.space_after = Pt(0)
        emit_inline(ap, " ".join(abstract), size=Pt(BODY_SIZE))

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ---------------------------------------------------------------------------
# block emitters
# ---------------------------------------------------------------------------
def emit_table(doc: Document, rows: list[list[str]], aligns: list[str], caption: str, number: int) -> None:
    header, body = rows[0], rows[1:]
    ncols = len(header)

    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_autofit(table)
    set_cell_margins(table)

    for i, text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        par = cell.paragraphs[0]
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.alignment = {
            "c": WD_ALIGN_PARAGRAPH.CENTER,
            "r": WD_ALIGN_PARAGRAPH.RIGHT,
        }.get(aligns[i], WD_ALIGN_PARAGRAPH.LEFT)
        run = par.add_run(strip_inline(text))
        run.bold = True
        run.font.name = TABLE_FONT
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(cell, HEADER_FILL)
    set_repeat_header(table.rows[0])

    for r, row in enumerate(body):
        cells = table.add_row().cells
        for i in range(ncols):
            text = row[i] if i < len(row) else ""
            cell = cells[i]
            cell.text = ""
            par = cell.paragraphs[0]
            par.paragraph_format.space_after = Pt(0)
            par.paragraph_format.alignment = {
                "c": WD_ALIGN_PARAGRAPH.CENTER,
                "r": WD_ALIGN_PARAGRAPH.RIGHT,
            }.get(aligns[i], WD_ALIGN_PARAGRAPH.LEFT)
            emit_inline(par, text, size=Pt(8.5))
            for run in par.runs:
                if run.font.name != MONO_FONT:
                    run.font.name = TABLE_FONT
            if r % 2 == 1:
                shade(cell, ZEBRA_FILL)

    cap = doc.add_paragraph(style="Caption")
    cr = cap.add_run(f"Table {number}. ")
    cr.bold = True
    cr.font.name = TABLE_FONT
    cr.font.size = Pt(8.5)
    cr.font.color.rgb = MUTED
    tr = cap.add_run(caption)
    tr.font.name = TABLE_FONT
    tr.font.size = Pt(8.5)
    tr.font.color.rgb = MUTED


def figure_number(path: Path) -> int:
    """Figure number comes from the filename, which is already in report order."""
    m = re.match(r"figure_(\d+)", path.name)
    if not m:
        raise ValueError(f"cannot read a figure number from {path.name}")
    return int(m.group(1))


def emit_figure(doc: Document, path: Path, caption: str, number: int) -> None:
    w, h = png_size(path)
    width_in = min(CONTENT_WIDTH_IN, 6.27)
    # Keep tall figures from running past a page; cap rendered height at 7.2in.
    if (width_in * h / w) > 7.2:
        width_in = 7.2 * w / h

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width_in))

    cap = doc.add_paragraph(style=FIGURE_CAPTION_STYLE)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"Figure {number}. ")
    cr.bold = True
    cr.font.name = TABLE_FONT
    cr.font.size = Pt(8.5)
    cr.font.color.rgb = MUTED
    tr = cap.add_run(caption)
    tr.font.name = TABLE_FONT
    tr.font.size = Pt(8.5)
    tr.font.color.rgb = MUTED


def emit_quote(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    shade(p, QUOTE_FILL)

    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "1F3B57")
    borders.append(left)
    pPr.append(borders)

    emit_inline(p, " ".join(lines), size=Pt(BODY_SIZE))


def emit_code(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        shade(p, CODE_FILL)
        run = p.add_run(line if line.strip() else " ")
        run.font.name = MONO_FONT
        run.font.size = Pt(9)
        run.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------------------
# main parse / render loop
# ---------------------------------------------------------------------------
def parse_table(lines: list[str]) -> tuple[list[list[str]], list[str]]:
    def split(row: str) -> list[str]:
        return [c.strip() for c in row.strip().strip("|").split("|")]

    header = split(lines[0])
    aligns = []
    for spec in split(lines[1]):
        if spec.startswith(":") and spec.endswith(":"):
            aligns.append("c")
        elif spec.endswith(":"):
            aligns.append("r")
        else:
            aligns.append("l")
    body = [split(r) for r in lines[2:]]
    return [header] + body, aligns


def main() -> int:
    if not SRC.exists():
        print(f"missing {SRC}", file=sys.stderr)
        return 1

    raw = SRC.read_text(encoding="utf-8").split("\n")

    doc = Document()
    page_setup(doc)
    build_styles(doc)

    # ---- front matter -----------------------------------------------------
    title, meta, abstract = "", [], []
    i = 0
    while i < len(raw) and not raw[i].startswith("## Table of contents"):
        s = raw[i].strip()
        if s.startswith("# "):
            title = s[2:].strip()
        elif s.startswith("**Course:**") or s.startswith("**Institution:**") \
                or s.startswith("**Date:**") or s.startswith("**Team"):
            meta.append(s)
        elif s and not s.startswith("#") and s != "---":
            abstract.append(s)
        i += 1
    title_page(doc, title, meta, abstract)

    # ---- table of contents ------------------------------------------------
    # Formatted directly rather than with the Heading 1 style, so the TOC field
    # does not list itself.
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(0)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = h.add_run("Table of contents")
    hr.font.name = HEAD_FONT
    hr.font.size = Pt(19)
    hr.font.color.rgb = ACCENT
    add_toc_field(doc.add_paragraph())
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---- list of figures, then list of tables ------------------------------
    def list_heading(text: str, first: bool) -> None:
        lh = doc.add_paragraph()
        lh.paragraph_format.space_before = Pt(0 if first else 22)
        lh.paragraph_format.space_after = Pt(6)
        lh.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lh.paragraph_format.keep_with_next = True
        lr = lh.add_run(text)
        lr.font.name = HEAD_FONT
        lr.font.size = Pt(19)
        lr.font.color.rgb = ACCENT

    list_heading("List of figures", first=True)
    add_toc_field(doc.add_paragraph(), TOC_FIGURES)

    list_heading("List of tables", first=False)
    add_toc_field(doc.add_paragraph(), TOC_TABLES)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # skip the markdown TOC list itself
    while i < len(raw) and not re.match(r"^## 0\.", raw[i]):
        i += 1

    table_no = fig_no = 0
    current_section = ""
    pending: list[str] = []
    first_h1 = True

    def flush() -> None:
        nonlocal pending
        if not pending:
            return
        p = doc.add_paragraph()
        emit_inline(p, " ".join(pending))
        pending = []

    def flush_figures(section: str) -> None:
        nonlocal fig_no
        for fname, caption in FIGURES.get(section, []):
            path = FIGDIR / fname
            if not path.exists():
                print(f"  ! missing figure {fname}", file=sys.stderr)
                continue
            fig_no += 1
            emit_figure(doc, path, caption, figure_number(path))

    while i < len(raw):
        line = raw[i]
        s = line.strip()

        if s.startswith("#"):
            flush()
            flush_figures(current_section)
            level = len(s) - len(s.lstrip("#"))
            text = strip_inline(s.lstrip("#").strip())
            current_section = text
            if level == 2:
                if not first_h1:
                    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                first_h1 = False
                doc.add_paragraph(text, style="Heading 1")
            else:
                doc.add_paragraph(text, style="Heading 2" if level == 3 else "Heading 3")
            i += 1
            continue

        if s.startswith("|"):
            flush()
            block = []
            while i < len(raw) and raw[i].strip().startswith("|"):
                block.append(raw[i].strip())
                i += 1
            if len(block) >= 2:
                rows, aligns = parse_table(block)
                table_no += 1
                caption = (TABLE_CAPTIONS[table_no - 1]
                           if table_no <= len(TABLE_CAPTIONS) else current_section)
                emit_table(doc, rows, aligns, caption, table_no)
            continue

        if s.startswith(">"):
            flush()
            block = []
            while i < len(raw) and raw[i].strip().startswith(">"):
                block.append(raw[i].strip().lstrip(">").strip())
                i += 1
            emit_quote(doc, block)
            continue

        if s.startswith("```"):
            flush()
            i += 1
            block = []
            while i < len(raw) and not raw[i].strip().startswith("```"):
                block.append(raw[i])
                i += 1
            i += 1
            emit_code(doc, block)
            continue

        m_ul = re.match(r"^- (.*)", s)
        m_ol = re.match(r"^(\d+)\. (.*)", s)
        if m_ul or m_ol:
            flush()
            text = m_ul.group(1) if m_ul else m_ol.group(2)
            i += 1
            # absorb wrapped continuation lines
            while i < len(raw):
                nxt = raw[i]
                if nxt.strip() and nxt.startswith((" ", "\t")) \
                        and not nxt.strip().startswith(("-", "|", ">")) \
                        and not re.match(r"^\s*\d+\. ", nxt):
                    text += " " + nxt.strip()
                    i += 1
                else:
                    break
            p = doc.add_paragraph(style="List Bullet" if m_ul else "List Number")
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            emit_inline(p, text)
            continue

        if s == "---" or not s:
            flush()
            i += 1
            continue

        pending.append(s)
        i += 1

    flush()
    flush_figures(current_section)

    force_field_update_on_open(doc)
    doc.save(OUT)

    expected = len(TABLE_CAPTIONS)
    print(f"wrote {OUT.relative_to(ROOT)}")
    print(f"  tables : {table_no} (captions defined: {expected})")
    print(f"  figures: {fig_no} of {len(list(FIGDIR.glob('*.png')))} available")
    if table_no != expected:
        print("  ! table count differs from curated captions — check TABLE_CAPTIONS",
              file=sys.stderr)
        return 2
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
